"""Text extraction for the AI-ingest path.

We extract a *truncated* representation of the document for the LLM. The goal
is metadata drafting, not document understanding: a few thousand chars from the
start, plus a few hundred from the end (to catch signature blocks, dates), is
plenty for tag inference.
"""

from __future__ import annotations

import io
import mimetypes
from dataclasses import dataclass
from pathlib import Path

# Lazy imports so the module loads cheaply when the AI path isn't used.


@dataclass
class ExtractResult:
    text: str
    mime: str
    truncated: bool
    note: str | None = None  # e.g. "scan, no extractable text"


_TEXT_PREFIXES = ("text/", "application/json", "application/xml")

# Extensions we treat as plain text even when mimetypes.guess_type() doesn't
# return a text/* MIME. The OS registry on Windows often has no entry for
# these formats, so without this fallback they'd fall through to the "no
# extractor" branch and (with vision enabled) trigger a confusing "sending 0
# rasterized pages" path.
_TEXT_EXTENSIONS = {
    ".toml", ".yaml", ".yml", ".ini", ".cfg", ".conf", ".log",
    ".env", ".md", ".markdown", ".rst", ".properties",
    ".csv", ".tsv",
}


def _guess_mime(path: Path) -> str:
    if path.suffix.lower() in _TEXT_EXTENSIONS:
        return "text/plain"
    mime, _ = mimetypes.guess_type(path.name)
    return mime or "application/octet-stream"


def _truncate(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    head = int(max_chars * 0.8)
    tail = max_chars - head - 40
    if tail < 200:
        return text[:max_chars], True
    return text[:head] + "\n\n[…truncated…]\n\n" + text[-tail:], True


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf not installed") from e
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)


def extract_text(path: Path, *, max_chars: int = 120_000) -> ExtractResult:
    mime = _guess_mime(path)

    if mime.startswith(_TEXT_PREFIXES):
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except OSError as e:
            return ExtractResult(text="", mime=mime, truncated=False, note=f"read error: {e}")
        out, truncated = _truncate(text, max_chars)
        return ExtractResult(text=out, mime=mime, truncated=truncated)

    if mime == "application/pdf":
        try:
            text = _extract_pdf(path)
        except Exception as e:
            return ExtractResult(text="", mime=mime, truncated=False, note=f"pdf extract failed: {e}")
        if not text.strip():
            return ExtractResult(
                text="",
                mime=mime,
                truncated=False,
                note="pdf has no extractable text (likely scanned; needs OCR or a vision model)",
            )
        out, truncated = _truncate(text, max_chars)
        return ExtractResult(text=out, mime=mime, truncated=truncated)

    if mime in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            text = _extract_docx(path)
        except Exception as e:
            return ExtractResult(text="", mime=mime, truncated=False, note=f"docx extract failed: {e}")
        out, truncated = _truncate(text, max_chars)
        return ExtractResult(text=out, mime=mime, truncated=truncated)

    if mime.startswith("image/"):
        return ExtractResult(
            text="",
            mime=mime,
            truncated=False,
            note="image — needs a vision-capable model",
        )

    return ExtractResult(
        text="",
        mime=mime,
        truncated=False,
        note=f"no extractor for mime={mime}",
    )


def extract_images(
    path: Path,
    *,
    max_pages: int = 3,
    max_dim: int = 1280,
) -> list[tuple[str, bytes]]:
    """Render a document to PNG bytes for vision LLMs.

    For PDFs: rasterize up to `max_pages` pages. Each page is scaled so its
    longest edge is `max_dim` px (PDFium uses a scale factor relative to 72dpi).
    For image files: read the original bytes (mime is whatever the file is).
    Other types return an empty list.

    Returns a list of (mime, bytes). Raises no errors for unsupported types —
    callers should treat an empty list as "nothing to send."
    """
    mime = _guess_mime(path)

    if mime == "application/pdf":
        try:
            import pypdfium2 as pdfium
        except ImportError:
            return []
        try:
            pdf = pdfium.PdfDocument(str(path))
        except Exception:
            return []
        out: list[tuple[str, bytes]] = []
        try:
            n = min(max_pages, len(pdf))
            for i in range(n):
                page = pdf[i]
                w_pt, h_pt = page.get_size()
                # PDFium points are 1/72 inch; scale = max_dim / longest_edge_pt
                longest = max(w_pt, h_pt) or 1.0
                scale = max(0.5, min(4.0, max_dim / longest))
                bitmap = page.render(scale=scale)
                pil = bitmap.to_pil()
                buf = io.BytesIO()
                pil.save(buf, format="PNG", optimize=True)
                out.append(("image/png", buf.getvalue()))
        finally:
            pdf.close()
        return out

    if mime.startswith("image/"):
        try:
            return [(mime, path.read_bytes())]
        except OSError:
            return []

    return []


__all__ = ["ExtractResult", "extract_text", "extract_images"]
